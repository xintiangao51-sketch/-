from docx import Document
import os

doc_path = r'C:\Users\高\Desktop\通衢隧道洞口坡面治理工程\通衢隧道_施工组织设计_劳务清包工版_20260529.docx'
output_path = r'C:\Users\高\Desktop\施工组织设计内容.md'

doc = Document(doc_path)
content = []

for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        content.append(text + '\n\n')

# 处理表格
for table in doc.tables:
    content.append('|' + '|'.join([cell.text.strip() for cell in table.rows[0].cells]) + '|\n')
    content.append('|' + '|'.join(['---'] * len(table.rows[0].cells)) + '|\n')
    for row in table.rows[1:]:
        content.append('|' + '|'.join([cell.text.strip() for cell in row.cells]) + '|\n')
    content.append('\n')

with open(output_path, 'w', encoding='utf-8-sig') as f:
    f.writelines(content)

print(f"施工组织设计内容已导出到：{output_path}")
