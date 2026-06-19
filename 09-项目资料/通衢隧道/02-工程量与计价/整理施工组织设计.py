from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 创建新文档
doc = Document()

# 设置默认字体
doc.styles['Normal'].font.name = u'微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
doc.styles['Normal'].font.size = Pt(12)

# 读取整理后的内容
content_path = r'C:\Users\高\Desktop\施工组织设计内容.md'
with open(content_path, 'r', encoding='utf-8-sig') as f:
    lines = f.readlines()

# 处理内容，生成docx
doc._current_table = None
for line in lines:
    line = line.strip()
    if not line:
        continue
    # 标题处理
    if line.startswith('# '):
        p = doc.add_heading(level=1)
        p.add_run(line[2:].strip())
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(18)
        p.runs[0].bold = True
        doc._current_table = None
    elif line.startswith('## '):
        p = doc.add_heading(level=2)
        p.add_run(line[3:].strip())
        p.runs[0].font.size = Pt(16)
        p.runs[0].bold = True
        doc._current_table = None
    elif line.startswith('### '):
        p = doc.add_heading(level=3)
        p.add_run(line[4:].strip())
        p.runs[0].font.size = Pt(14)
        p.runs[0].bold = True
        doc._current_table = None
    elif line.startswith('■ '):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(line[2:].strip())
        doc._current_table = None
    elif line.startswith('|') and line.count('|') > 2:
        # 处理表格
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if doc._current_table is None:
            # 表头
            doc._current_table = doc.add_table(rows=1, cols=len(cells))
            doc._current_table.style = 'Table Grid'
            for i, cell in enumerate(cells):
                doc._current_table.cell(0, i).text = cell
                doc._current_table.cell(0, i).paragraphs[0].runs[0].bold = True
        elif all(c.strip() == '-' or c.strip() == '' for c in line.split('|') if c.strip()):
            # 分隔行，跳过
            continue
        else:
            # 数据行
            if len(cells) == len(doc._current_table.columns):
                row = doc._current_table.add_row()
                for i, cell in enumerate(cells):
                    row.cells[i].text = cell
    else:
        # 普通段落
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc._current_table = None

# 调整表格宽度
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1  # 垂直居中
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True

# 保存文档
output_path = r'C:\Users\高\Desktop\通衢隧道_施工组织设计_劳务清包工版_整理版_20260530.docx'
doc.save(output_path)
print(f"整理后的施工组织设计已保存到：{output_path}")
