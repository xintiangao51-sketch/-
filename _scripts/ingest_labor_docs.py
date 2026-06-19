from pathlib import Path
import json, os, shutil, sys, traceback, re
from datetime import datetime

ROOT = Path(r'D:/知识库')
DESKTOP = Path(r'C:/Users/20752/Desktop')
FILES = [
    '2026年市政工程劳务指导单价222.xlsx',
    '建筑行业税务征收办法.docx',
    '劳务扩大报价分析提示词.md',
    '企业级对下劳务分包单价汇总表2020xls.xls',
    '鄯善启创劳务扩大报价分析_核查表.xlsx',
    '鄯善启创劳务扩大报价分析报告.docx',
    '鄯善启创劳务清包工报价表.xlsx',
    '鄯善启创劳务清包工协议.docx',
    '田让存旋挖钻两项目最终结算单.xlsx',
]
OUT = ROOT / '08_项目案例库' / '劳务分包与扩大劳务报价' / '2026-劳务价格与鄯善启创案例'
SRC = OUT / '00_原始附件'
TXT = OUT / '01_文本抽取'
TAB = OUT / '02_表格抽取'
LOG = OUT / 'ingest_log.json'

def find_source(filename):
    """在多个路径中查找源文件"""
    candidates = [
        ROOT / filename,               # D:\知识库\
        DESKTOP / filename,            # 桌面
        SRC / filename,                # 已归档
    ]
    for p in candidates:
        if p.exists():
            return p
    return None
for p in [SRC, TXT, TAB]:
    p.mkdir(parents=True, exist_ok=True)

meta = {'run_time': datetime.now().isoformat(timespec='seconds'), 'root': str(ROOT), 'files': []}

def clean_cell(v):
    if v is None:
        return ''
    s = str(v)
    s = s.replace('\r\n','\n').replace('\r','\n')
    return s.strip()

def md_escape_table(s):
    return clean_cell(s).replace('|','\\|').replace('\n','<br>')

def extract_xlsx(path: Path, outbase: str):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=False, read_only=False)
    info = {'type':'xlsx', 'sheets': []}
    for ws in wb.worksheets:
        max_row, max_col = ws.max_row, ws.max_column
        # count non-empty rows/cells and identify used bounds
        rows = []
        nonempty = 0
        for row in ws.iter_rows(values_only=True):
            vals = [clean_cell(v) for v in row]
            if any(vals):
                nonempty += 1
                rows.append(vals)
        sheet_info = {'name': ws.title, 'max_row': max_row, 'max_col': max_col, 'nonempty_rows': nonempty}
        info['sheets'].append(sheet_info)
        # write TSV full content
        safe = re.sub(r'[\\/:*?"<>|\s]+','_', ws.title)[:80]
        tsv = TAB / f'{outbase}__{safe}.tsv'
        with open(tsv, 'w', encoding='utf-8-sig', newline='') as f:
            for vals in rows:
                f.write('\t'.join(vals).replace('\x00','') + '\n')
        # write markdown preview first 80 non-empty rows, 20 columns
        md = TXT / f'{outbase}__{safe}_preview.md'
        max_preview_rows = min(len(rows), 120)
        max_preview_cols = min(max((len(r) for r in rows), default=0), 20)
        with open(md, 'w', encoding='utf-8') as f:
            f.write(f'# {path.name} / {ws.title}\n\n')
            f.write(f'- used_rows: {max_row}\n- used_cols: {max_col}\n- nonempty_rows: {nonempty}\n\n')
            if max_preview_rows and max_preview_cols:
                f.write('| ' + ' | '.join([f'C{i+1}' for i in range(max_preview_cols)]) + ' |\n')
                f.write('| ' + ' | '.join(['---']*max_preview_cols) + ' |\n')
                for vals in rows[:max_preview_rows]:
                    vals = (vals + ['']*max_preview_cols)[:max_preview_cols]
                    f.write('| ' + ' | '.join(md_escape_table(v) for v in vals) + ' |\n')
    return info

def extract_xls(path: Path, outbase: str):
    import xlrd
    # xlrd ignores sheet protection; if file-level encrypted, this will fail.
    wb = xlrd.open_workbook(str(path), formatting_info=False, on_demand=True)
    info = {'type':'xls', 'sheets': []}
    for sh in wb.sheets():
        sheet_info = {'name': sh.name, 'max_row': sh.nrows, 'max_col': sh.ncols, 'nonempty_rows': 0}
        rows = []
        for r in range(sh.nrows):
            vals = [clean_cell(sh.cell_value(r,c)) for c in range(sh.ncols)]
            if any(vals):
                sheet_info['nonempty_rows'] += 1
                rows.append(vals)
        info['sheets'].append(sheet_info)
        safe = re.sub(r'[\\/:*?"<>|\s]+','_', sh.name)[:80]
        tsv = TAB / f'{outbase}__{safe}.tsv'
        with open(tsv, 'w', encoding='utf-8-sig', newline='') as f:
            for vals in rows:
                f.write('\t'.join(vals).replace('\x00','') + '\n')
        md = TXT / f'{outbase}__{safe}_preview.md'
        max_preview_rows = min(len(rows), 160)
        max_preview_cols = min(max((len(r) for r in rows), default=0), 20)
        with open(md, 'w', encoding='utf-8') as f:
            f.write(f'# {path.name} / {sh.name}\n\n')
            f.write(f'- used_rows: {sh.nrows}\n- used_cols: {sh.ncols}\n- nonempty_rows: {sheet_info["nonempty_rows"]}\n- password_note: 用户提供密码=2020；本次xlrd可直接读取，说明可能为工作簿/工作表保护或弱加密，非文件级阻断加密。\n\n')
            if max_preview_rows and max_preview_cols:
                f.write('| ' + ' | '.join([f'C{i+1}' for i in range(max_preview_cols)]) + ' |\n')
                f.write('| ' + ' | '.join(['---']*max_preview_cols) + ' |\n')
                for vals in rows[:max_preview_rows]:
                    vals = (vals + ['']*max_preview_cols)[:max_preview_cols]
                    f.write('| ' + ' | '.join(md_escape_table(v) for v in vals) + ' |\n')
    wb.release_resources()
    return info

def extract_docx(path: Path, outbase: str):
    from docx import Document
    doc = Document(str(path))
    lines = []
    for p in doc.paragraphs:
        text = clean_cell(p.text)
        if text:
            lines.append(text)
    table_count = 0
    for ti, table in enumerate(doc.tables, start=1):
        table_count += 1
        lines.append(f'\n[表格 {ti}]')
        for row in table.rows:
            vals = [clean_cell(cell.text) for cell in row.cells]
            if any(vals):
                lines.append('\t'.join(vals))
    out = TXT / f'{outbase}.md'
    with open(out, 'w', encoding='utf-8') as f:
        f.write(f'# {path.name}\n\n')
        f.write('\n\n'.join(lines))
    return {'type':'docx', 'paragraphs': len(doc.paragraphs), 'tables': table_count, 'nonempty_blocks': len(lines)}

def extract_md(path: Path, outbase: str):
    data = path.read_bytes()
    text = None
    for enc in ['utf-8-sig','utf-8','gbk','gb2312']:
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            pass
    if text is None:
        text = data.decode('utf-8', errors='ignore')
    out = TXT / f'{outbase}.md'
    out.write_text(text, encoding='utf-8')
    return {'type':'md', 'chars': len(text), 'lines': text.count('\n')+1}

for fname in FILES:
    src = find_source(fname)
    if src is None:
        rec = {'file': fname, 'source': 'NOT FOUND', 'exists': False, 'searched': [str(ROOT/fname), str(DESKTOP/fname), str(SRC/fname)]}
        meta['files'].append(rec)
        continue
    rec = {'file': fname, 'source': str(src), 'exists': True}
    dst = SRC / fname
    shutil.copy2(src, dst)
    rec['archived_to'] = str(dst)
    rec['size_bytes'] = dst.stat().st_size
    outbase = Path(fname).stem
    try:
        suf = src.suffix.lower()
        if suf == '.xlsx' or suf == '.xlsm':
            rec['extract'] = extract_xlsx(dst, outbase)
        elif suf == '.xls':
            rec['extract'] = extract_xls(dst, outbase)
        elif suf == '.docx':
            rec['extract'] = extract_docx(dst, outbase)
        elif suf in ['.md','.txt']:
            rec['extract'] = extract_md(dst, outbase)
        else:
            rec['extract'] = {'type':'unsupported'}
        rec['status'] = 'ok'
    except Exception as e:
        rec['status'] = 'error'
        rec['error'] = repr(e)
        rec['traceback'] = traceback.format_exc()
    meta['files'].append(rec)
LOG.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding='utf-8')
print(json.dumps(meta, ensure_ascii=False, indent=2))
